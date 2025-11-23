# -*- coding: utf-8 -*-
"""
DOCX Handler - Gestion des documents Word avec préservation de la mise en forme
Auteur : Renaud LOISON
"""

from typing import List, Tuple, Optional
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from copy import deepcopy


class DocxTranslationHandler:
    """Gère la traduction de documents Word en préservant la mise en forme"""

    def __init__(self):
        self.doc = None

    def load_document(self, file_path: str) -> Document:
        """Charge un document Word"""
        self.doc = Document(file_path)
        return self.doc

    def extract_text_with_metadata(self) -> List[Tuple[str, dict]]:
        """
        Extrait le texte avec métadonnées de mise en forme
        Retourne: [(texte, metadata), ...]
        """
        text_items = []

        # Parcourir tous les paragraphes
        for para_idx, paragraph in enumerate(self.doc.paragraphs):
            if not paragraph.text.strip():
                # Conserver les paragraphes vides pour la pagination
                text_items.append(("", {
                    'type': 'paragraph',
                    'para_idx': para_idx,
                    'alignment': paragraph.alignment,
                    'style': paragraph.style.name if paragraph.style else None,
                    'runs': []
                }))
                continue

            # Extraire les runs (portions de texte avec style uniforme)
            runs_metadata = []
            full_text = ""

            for run in paragraph.runs:
                if run.text.strip():
                    runs_metadata.append({
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'font_color': self._get_color(run.font.color),
                        'highlight_color': run.font.highlight_color,
                    })
                    full_text += run.text

            if full_text.strip():
                text_items.append((full_text, {
                    'type': 'paragraph',
                    'para_idx': para_idx,
                    'alignment': paragraph.alignment,
                    'style': paragraph.style.name if paragraph.style else None,
                    'runs': runs_metadata
                }))

        # Parcourir les tableaux
        for table_idx, table in enumerate(self.doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        if para.text.strip():
                            runs_metadata = []
                            full_text = ""

                            for run in para.runs:
                                if run.text.strip():
                                    runs_metadata.append({
                                        'text': run.text,
                                        'bold': run.bold,
                                        'italic': run.italic,
                                        'underline': run.underline,
                                        'font_name': run.font.name,
                                        'font_size': run.font.size,
                                        'font_color': self._get_color(run.font.color),
                                    })
                                    full_text += run.text

                            if full_text.strip():
                                text_items.append((full_text, {
                                    'type': 'table_cell',
                                    'table_idx': table_idx,
                                    'row_idx': row_idx,
                                    'cell_idx': cell_idx,
                                    'para_idx': cell.paragraphs.index(para),
                                    'runs': runs_metadata
                                }))

        return text_items

    def _get_color(self, color_obj):
        """Extrait la couleur RGB"""
        if color_obj and color_obj.rgb:
            return color_obj.rgb
        return None

    def apply_translation(self, translations: List[Tuple[str, dict]], output_path: str):
        """
        Applique les traductions au document en préservant la mise en forme
        translations: [(texte_traduit, metadata), ...]
        """
        para_translations = {}
        table_translations = {}

        # Organiser les traductions
        for translated_text, metadata in translations:
            if metadata['type'] == 'paragraph':
                para_translations[metadata['para_idx']] = (translated_text, metadata)
            elif metadata['type'] == 'table_cell':
                key = (metadata['table_idx'], metadata['row_idx'],
                       metadata['cell_idx'], metadata['para_idx'])
                table_translations[key] = (translated_text, metadata)

        # Appliquer aux paragraphes
        for para_idx, paragraph in enumerate(self.doc.paragraphs):
            if para_idx in para_translations:
                translated_text, metadata = para_translations[para_idx]
                self._apply_paragraph_translation(paragraph, translated_text, metadata)

        # Appliquer aux tableaux
        for table_idx, table in enumerate(self.doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        key = (table_idx, row_idx, cell_idx, para_idx)
                        if key in table_translations:
                            translated_text, metadata = table_translations[key]
                            self._apply_paragraph_translation(para, translated_text, metadata)

        # Sauvegarder
        self.doc.save(output_path)

    def _apply_paragraph_translation(self, paragraph, translated_text: str, metadata: dict):
        """Applique la traduction à un paragraphe en préservant le style"""
        if not translated_text.strip():
            return

        # Vider le paragraphe
        for run in paragraph.runs:
            run.text = ""

        # Si on a des métadonnées de runs, essayer de préserver les styles
        runs_meta = metadata.get('runs', [])

        if runs_meta and len(runs_meta) == 1:
            # Un seul run original, appliquer le même style
            run = paragraph.add_run(translated_text)
            self._apply_run_style(run, runs_meta[0])

        elif runs_meta and len(runs_meta) > 1:
            # Plusieurs runs, essayer de distribuer intelligemment
            # Pour l'instant, on applique le style du premier run à tout le texte
            # TODO: améliorer la distribution des styles
            run = paragraph.add_run(translated_text)
            self._apply_run_style(run, runs_meta[0])

        else:
            # Pas de métadonnées, texte simple
            paragraph.add_run(translated_text)

        # Appliquer l'alignement du paragraphe
        if metadata.get('alignment'):
            paragraph.alignment = metadata['alignment']

    def _apply_run_style(self, run, style_meta: dict):
        """Applique le style d'un run"""
        if style_meta.get('bold') is not None:
            run.bold = style_meta['bold']
        if style_meta.get('italic') is not None:
            run.italic = style_meta['italic']
        if style_meta.get('underline') is not None:
            run.underline = style_meta['underline']
        if style_meta.get('font_name'):
            run.font.name = style_meta['font_name']
        if style_meta.get('font_size'):
            run.font.size = style_meta['font_size']
        if style_meta.get('font_color'):
            run.font.color.rgb = style_meta['font_color']
        if style_meta.get('highlight_color'):
            run.font.highlight_color = style_meta['highlight_color']

    def get_statistics(self) -> dict:
        """Retourne des statistiques sur le document"""
        if not self.doc:
            return {}

        return {
            'paragraphs': len(self.doc.paragraphs),
            'tables': len(self.doc.tables),
            'sections': len(self.doc.sections),
            'total_chars': sum(len(p.text) for p in self.doc.paragraphs),
        }


class DocxProcessor:
    """Processeur de documents Word pour traduction"""

    @staticmethod
    def extract_texts_for_translation(file_path: str) -> Tuple[List[str], List[dict], DocxTranslationHandler]:
        """
        Extrait les textes à traduire d'un document Word
        Retourne: (liste_textes, liste_metadonnées, handler)
        """
        handler = DocxTranslationHandler()
        handler.load_document(file_path)

        text_items = handler.extract_text_with_metadata()

        texts = [text for text, _ in text_items]
        metadata = [meta for _, meta in text_items]

        return texts, metadata, handler

    @staticmethod
    def apply_translations(handler: DocxTranslationHandler, texts: List[str],
                          metadata: List[dict], output_path: str):
        """
        Applique les traductions au document
        """
        translations = list(zip(texts, metadata))
        handler.apply_translation(translations, output_path)
